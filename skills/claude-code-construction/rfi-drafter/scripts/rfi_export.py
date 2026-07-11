#!/usr/bin/env python3
"""
rfi_export.py — Generate an RFI Word document (.docx).

Two modes:
  1. Template mode: populate the user's RFI template with field data
  2. Generic mode: build a standard RFI document from scratch (fallback)

Template mode (preferred):
  python rfi_export.py \
    --template firm_rfi_form.docx \
    --mapping rfi_template_map.json \
    --data rfi_draft.json \
    --output RFI-026.docx

Generic mode (no template):
  python rfi_export.py \
    --data rfi_draft.json \
    --output RFI-026.docx

Input JSON format (--data):
{
  "project_name": "...",
  "project_number": "...",
  "rfi_number": "RFI-026",
  "date": "2026-04-05",
  "from": { "name": "...", "company": "..." },
  "to": { "name": "...", "company": "..." },
  "subject": "...",
  "spec_section": "...",
  "drawing_ref": "...",
  "description": "...",
  "suggested_resolution": "...",
  "impact": "...",
  "attachments": ["..."]
}

Mapping JSON format (--mapping):
{
  "template_path": "firm_rfi_form.docx",
  "template_hash": "sha256...",
  "field_mappings": [
    {"field": "project_name", "type": "table_cell", "table_index": 0, "row": 0, "col": 1},
    {"field": "description", "type": "placeholder", "search_text": "[DESCRIPTION]"},
    {"field": "subject", "type": "content_control", "tag": "Subject"}
  ]
}

Requires: python-docx (pip install python-docx)

This is a THIN FORMATTER. It renders pre-populated content into the
document. All content decisions are made upstream by Claude Code.
"""

import argparse
import hashlib
import json
import sys
from datetime import date
from pathlib import Path

from shared import safe_output_path

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


# ---------------------------------------------------------------------------
# Template population (Mode 1)
# ---------------------------------------------------------------------------

def _resolve_value(rfi_data: dict, field: str) -> str:
    """Resolve a field name to its value from the RFI data dict.

    Handles nested keys like 'from.name' and 'from.company'.
    """
    if "." in field:
        parts = field.split(".", 1)
        sub = rfi_data.get(parts[0], {})
        if isinstance(sub, dict):
            return str(sub.get(parts[1], ""))
        return ""

    val = rfi_data.get(field, "")
    if isinstance(val, list):
        return "\n".join(f"{i+1}. {item}" for i, item in enumerate(val))
    return str(val)


def _set_cell_text_preserve_format(cell, text: str):
    """Replace cell text while preserving the first run's formatting."""
    para = cell.paragraphs[0]
    if para.runs:
        for run in para.runs[1:]:
            run.text = ""
        para.runs[0].text = text
    else:
        para.text = text


def _replace_placeholder_in_paragraphs(paragraphs, search_text: str, value: str):
    """Find a placeholder string in paragraphs and replace with value."""
    for para in paragraphs:
        full_text = para.text
        if search_text not in full_text:
            continue
        # Replace across runs
        if para.runs:
            combined = "".join(r.text for r in para.runs)
            if search_text in combined:
                new_text = combined.replace(search_text, value)
                para.runs[0].text = new_text
                for run in para.runs[1:]:
                    run.text = ""
        else:
            para.text = full_text.replace(search_text, value)
        return True
    return False


def populate_template(template_path: str, mapping: dict, rfi_data: dict,
                      output_path: str):
    """Populate a .docx template using the field mapping."""
    if not HAS_DOCX:
        print("python-docx not installed. Install: pip install python-docx",
              file=sys.stderr)
        sys.exit(1)

    doc = Document(template_path)

    for entry in mapping.get("field_mappings", []):
        field = entry["field"]
        value = _resolve_value(rfi_data, field)
        map_type = entry["type"]

        if map_type == "table_cell":
            tbl_idx = entry["table_index"]
            row_idx = entry["row"]
            col_idx = entry["col"]
            if tbl_idx < len(doc.tables):
                table = doc.tables[tbl_idx]
                if row_idx < len(table.rows) and col_idx < len(table.columns):
                    cell = table.cell(row_idx, col_idx)
                    _set_cell_text_preserve_format(cell, value)

        elif map_type == "placeholder":
            search = entry["search_text"]
            # Search paragraphs in body
            found = _replace_placeholder_in_paragraphs(doc.paragraphs,
                                                       search, value)
            # Also search table cells
            if not found:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if _replace_placeholder_in_paragraphs(
                                    cell.paragraphs, search, value):
                                found = True
                                break
                        if found:
                            break
                    if found:
                        break

        elif map_type == "content_control":
            # python-docx doesn't natively support content controls,
            # but we can search the underlying XML for tagged SDTs
            tag = entry.get("tag", "")
            _fill_content_control(doc, tag, value)

    doc.save(output_path)
    return output_path


def _fill_content_control(doc, tag: str, value: str):
    """Fill a structured document tag (content control) by its tag name.

    Walks the document XML to find <w:sdt> elements with matching
    <w:tag w:val="..."/> and replaces the text content.
    """
    from docx.oxml.ns import qn

    for sdt in doc.element.iter(qn("w:sdt")):
        pr = sdt.find(qn("w:sdtPr"))
        if pr is None:
            continue
        tag_elem = pr.find(qn("w:tag"))
        if tag_elem is not None and tag_elem.get(qn("w:val")) == tag:
            # Find the content element and replace text
            content = sdt.find(qn("w:sdtContent"))
            if content is not None:
                for p in content.iter(qn("w:p")):
                    for r in p.iter(qn("w:r")):
                        for t in r.iter(qn("w:t")):
                            t.text = value
                            return


def hash_template(template_path: str) -> str:
    """SHA-256 hash of the template file for change detection."""
    h = hashlib.sha256()
    with open(template_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


# ---------------------------------------------------------------------------
# Generic document generation (Mode 2 — fallback when no template)
# ---------------------------------------------------------------------------

def generate_generic_docx(rfi_data: dict, output_path: str):
    """Generate a standard RFI Word document from scratch."""
    if not HAS_DOCX:
        print("python-docx not installed. Install: pip install python-docx",
              file=sys.stderr)
        sys.exit(1)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("REQUEST FOR INFORMATION", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Header fields table
    header_table = doc.add_table(rows=6, cols=2)
    header_table.style = "Table Grid"

    fields = [
        ("Project:", rfi_data.get("project_name", "")),
        ("Project No:", rfi_data.get("project_number", "")),
        ("RFI No:", rfi_data.get("rfi_number", "TBD")),
        ("Date:", rfi_data.get("date", date.today().isoformat())),
        ("From:",
         f"{rfi_data['from']['company']} — {rfi_data['from']['name']}"),
        ("To:",
         f"{rfi_data['to']['company']} — {rfi_data['to']['name']}"),
    ]

    for i, (label, value) in enumerate(fields):
        cell_label = header_table.cell(i, 0)
        cell_value = header_table.cell(i, 1)
        cell_label.text = label
        cell_value.text = value
        for paragraph in cell_label.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph()

    # Subject / Spec / Drawing Ref
    for label, key, default in [
        ("Subject: ", "subject", ""),
        ("Spec Section: ", "spec_section", "N/A"),
        ("Drawing Ref: ", "drawing_ref", ""),
    ]:
        para = doc.add_paragraph()
        run = para.add_run(label)
        run.bold = True
        para.add_run(rfi_data.get(key, default))

    doc.add_paragraph()

    # Content sections
    for heading, key in [
        ("DESCRIPTION", "description"),
        ("SUGGESTED RESOLUTION", "suggested_resolution"),
        ("IMPACT IF NOT RESOLVED", "impact"),
    ]:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(rfi_data.get(key, ""))

    # Attachments
    attachments = rfi_data.get("attachments", [])
    if attachments:
        doc.add_heading("ATTACHMENTS", level=2)
        for i, attachment in enumerate(attachments, 1):
            doc.add_paragraph(f"{i}. {attachment}")

    # Response section (blank for A/E)
    doc.add_paragraph()
    doc.add_paragraph("_" * 60)
    doc.add_heading("RESPONSE (to be completed by design professional)",
                    level=2)
    doc.add_paragraph()
    doc.add_paragraph("Response:")
    doc.add_paragraph()
    doc.add_paragraph()

    resp_table = doc.add_table(rows=3, cols=2)
    resp_table.style = "Table Grid"
    for i, (label, value) in enumerate([
        ("Responded By:", ""),
        ("Date:", ""),
        ("Cost Impact:", "\u2610 Yes  \u2610 No  \u2610 TBD"),
    ]):
        resp_table.cell(i, 0).text = label
        resp_table.cell(i, 1).text = value

    doc.save(output_path)
    return output_path


# ---------------------------------------------------------------------------
# Plaintext fallback
# ---------------------------------------------------------------------------

def generate_plaintext(rfi_data: dict) -> str:
    """Generate plaintext RFI for environments without python-docx."""
    lines = [
        "REQUEST FOR INFORMATION",
        "=" * 50,
        "",
        f"Project:      {rfi_data.get('project_name', '')}",
        f"Project No:   {rfi_data.get('project_number', '')}",
        f"RFI No:       {rfi_data.get('rfi_number', 'TBD')}",
        f"Date:         {rfi_data.get('date', '')}",
        f"From:         {rfi_data['from']['company']} — "
        f"{rfi_data['from']['name']}",
        f"To:           {rfi_data['to']['company']} — "
        f"{rfi_data['to']['name']}",
        "",
        f"Subject:      {rfi_data.get('subject', '')}",
        f"Spec Section: {rfi_data.get('spec_section', 'N/A')}",
        f"Drawing Ref:  {rfi_data.get('drawing_ref', '')}",
        "",
        "DESCRIPTION:",
        rfi_data.get("description", ""),
        "",
        "SUGGESTED RESOLUTION:",
        rfi_data.get("suggested_resolution", ""),
        "",
        "IMPACT IF NOT RESOLVED:",
        rfi_data.get("impact", ""),
        "",
    ]
    attachments = rfi_data.get("attachments", [])
    if attachments:
        lines.append("ATTACHMENTS:")
        for i, a in enumerate(attachments, 1):
            lines.append(f"  {i}. {a}")
    return "\n".join(lines)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Generate RFI document (template or generic)")
    parser.add_argument("--data", dest="input",
                        help="JSON file with RFI field data")
    parser.add_argument("--template",
                        help="Path to firm's .docx RFI template")
    parser.add_argument("--mapping",
                        help="JSON file with template field mapping")
    parser.add_argument("--output", required=True,
                        help="Output file path (.docx)")

    args = parser.parse_args()

    # Load RFI data
    if args.input:
        with open(args.input) as f:
            rfi_data = json.load(f)
    else:
        print("ERROR: --data is required", file=sys.stderr)
        sys.exit(1)

    output = safe_output_path(args.output)

    if args.template and args.mapping:
        # Template mode: populate user's template
        with open(args.mapping) as f:
            mapping = json.load(f)
        populate_template(args.template, mapping, rfi_data, str(output))
        print(f"RFI populated from template → {output}", file=sys.stderr)

    elif output.suffix == ".docx" or str(output).endswith(".docx"):
        # Generic mode: build from scratch
        generate_generic_docx(rfi_data, str(output))
        print(f"RFI document (generic) → {output}", file=sys.stderr)

    else:
        # Plaintext fallback
        text = generate_plaintext(rfi_data)
        with open(output, "w") as f:
            f.write(text)
        print(f"RFI plaintext → {output}", file=sys.stderr)

    # JSON to stdout for pipeline use
    print(json.dumps(rfi_data, indent=2))


if __name__ == "__main__":
    main()