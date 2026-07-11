"""Block renderers for subcontract document formatting.

This module is a THIN FORMATTER — it renders pre-populated content into
python-docx elements. It makes no content decisions. All text, tables,
and lists arrive fully populated from Claude Code via template_data.json.

Block types: table, bullet_list, numbered_list, info_table
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Currency utilities (used by Claude Code for validation + formatter fallback)
# ---------------------------------------------------------------------------

_ONES = [
    "", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight",
    "Nine", "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen",
    "Sixteen", "Seventeen", "Eighteen", "Nineteen",
]
_TENS = [
    "", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy",
    "Eighty", "Ninety",
]


def _two_digits(n):
    if n < 20:
        return _ONES[n]
    return (_TENS[n // 10] + ("-" + _ONES[n % 10] if n % 10 else "")).strip()


def _three_digits(n):
    if n == 0:
        return ""
    if n >= 100:
        return f"{_ONES[n // 100]} Hundred" + (
            f" {_two_digits(n % 100)}" if n % 100 else ""
        )
    return _two_digits(n)


def amount_to_words(amount):
    """Convert integer dollar amount to written words."""
    amt = int(round(amount))
    if amt == 0:
        return "Zero Dollars and No Cents"
    parts = []
    for divisor, label in [(1_000_000_000, "Billion"), (1_000_000, "Million"),
                           (1_000, "Thousand")]:
        chunk = amt // divisor
        if chunk:
            parts.append(f"{_three_digits(chunk)} {label}")
        amt %= divisor
    if amt:
        parts.append(_three_digits(amt))
    return " ".join(parts) + " Dollars and No Cents"


def format_currency(amount):
    """Format number as $X,XXX,XXX.XX."""
    return f"${amount:,.2f}"


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def render_block(doc, block):
    """Dispatch a block dict to the appropriate renderer."""
    block_type = block.get("type", "")
    renderer = _RENDERERS.get(block_type)
    if renderer:
        renderer(doc, block)
    else:
        print(f"WARNING: Unknown block type '{block_type}', skipping")


def render_table(doc, block):
    """Render a table block: headers + rows of string values."""
    label = block.get("label", "")
    headers = block.get("headers", [])
    rows = block.get("rows", [])

    if label:
        doc.add_paragraph()
        doc.add_paragraph(label).runs[0].bold = True

    if not headers or not rows:
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            if c_idx < len(table.columns):
                table.rows[r_idx + 1].cells[c_idx].text = str(val) if val else ""


def render_bullet_list(doc, block):
    """Render a labeled bullet list."""
    label = block.get("label", "")
    items = block.get("items", [])

    if label:
        doc.add_paragraph()
        doc.add_paragraph(label).runs[0].bold = True

    for item in items:
        doc.add_paragraph(str(item), style="List Bullet")


def render_numbered_list(doc, block):
    """Render a labeled numbered list."""
    label = block.get("label", "")
    items = block.get("items", [])

    if label:
        doc.add_paragraph()
        doc.add_paragraph(label).runs[0].bold = True

    for i, item in enumerate(items, 1):
        doc.add_paragraph(f"{i}. {item}")


def render_info_table(doc, block):
    """Render a two-column key-value info table (cover page style)."""
    rows = block.get("rows", [])
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        label_cell = table.rows[r_idx].cells[0]
        value_cell = table.rows[r_idx].cells[1]
        label_cell.text = str(row.get("label", ""))
        value_cell.text = str(row.get("value", ""))
        # Bold the label column
        for p in label_cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
        for p in value_cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Signature block (structural formatting only, no content decisions)
# ---------------------------------------------------------------------------

def write_signature_block(doc, scope):
    """Write execution signature blocks for both parties."""
    contractor = scope.get("contractor_name", "[CONTRACTOR]")
    sub_name = scope.get("subcontractor", {}).get("company_name", "[SUBCONTRACTOR]")

    doc.add_page_break()
    doc.add_heading("Execution", level=1)

    doc.add_paragraph(
        "IN WITNESS WHEREOF, the parties have executed this Subcontract Agreement "
        "as of the date first written above."
    )
    doc.add_paragraph()

    for party in [contractor, sub_name]:
        p = doc.add_paragraph()
        p.add_run(party).bold = True
        doc.add_paragraph()
        doc.add_paragraph(f"By: {'_' * 40}")
        doc.add_paragraph(f"Name: {'_' * 37}")
        doc.add_paragraph(f"Title: {'_' * 37}")
        doc.add_paragraph(f"Date: {'_' * 38}")
        doc.add_paragraph()


# ---------------------------------------------------------------------------
# Header / footer (structural formatting only, no content decisions)
# ---------------------------------------------------------------------------

def write_header(doc, contractor_name, subcontract_number):
    """Add right-aligned header: Contractor | Subcontract Number."""
    from docx.oxml.ns import qn

    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.clear()
        run = p.add_run(f"{contractor_name}  |  {subcontract_number}")
        run.font.size = Pt(8)
        run.font.name = "Arial"


def write_footer(doc, project_name):
    """Add footer: Project name | CONFIDENTIAL | Page N."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.clear()

        # Left portion: project name
        run = p.add_run(f"{project_name}  |  CONFIDENTIAL  |  Page ")
        run.font.size = Pt(8)
        run.font.name = "Arial"

        # Auto page number field
        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")
        instr_text.set(qn("xml:space"), "preserve")
        instr_text.text = " PAGE "
        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")

        run2 = p.add_run()
        run2.font.size = Pt(8)
        run2.font.name = "Arial"
        run2._r.append(fld_char_begin)
        run2._r.append(instr_text)
        run2._r.append(fld_char_end)


# Dispatch table — keyed on block type, not content type
_RENDERERS = {
    "table": render_table,
    "bullet_list": render_bullet_list,
    "numbered_list": render_numbered_list,
    "info_table": render_info_table,
}
