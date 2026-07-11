#!/usr/bin/env python3
"""Generate a subcontract Word document from fully-populated JSON data.

This is a THIN FORMATTER. It renders pre-populated content from
template_data.json into a .docx file. It makes no content decisions —
all article text, tables, and lists are generated upstream by Claude Code.
"""

import argparse
import json
import sys
from pathlib import Path

from shared import safe_output_path

from subcontract_formatters import render_block, write_signature_block, write_header, write_footer


def generate_subcontract(template_data_path, scope_data_path, output="Subcontract.docx"):
    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)

    with open(template_data_path) as f:
        template = json.load(f)
    with open(scope_data_path) as f:
        scope = json.load(f)

    # Use existing .docx template if provided
    docx_template = template.get("docx_template_path")
    if docx_template and Path(docx_template).exists():
        doc = Document(docx_template)
    else:
        doc = Document()

    _format_document(doc, template, scope)

    out = safe_output_path(output)
    doc.save(str(out))
    print(f"OK: {out}")


def _format_document(doc, template, scope):
    """Render all pre-populated content into the document."""
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Document setup
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Header and footer
    contractor_name = scope.get("contractor_name", "")
    sc_number = scope.get("subcontract_number", "")
    project_name = scope.get("project_name", "")
    if contractor_name or sc_number:
        write_header(doc, contractor_name, sc_number)
    if project_name:
        write_footer(doc, project_name)

    # Cover page
    cover = template.get("cover_page", {})
    _render_section(doc, cover, is_cover=True)

    # Articles
    for article in template.get("articles", []):
        art_num = article.get("number", "")
        art_title = article.get("title", "")

        heading = doc.add_heading(f"Article {art_num} — {art_title}", level=1)
        heading.runs[0].font.size = Pt(12)

        text = article.get("text", "")
        if not text:
            print(f"WARNING: Article {art_num} ({art_title}) has empty text")
            text = f"[Article {art_num} content not provided]"

        # Render text as paragraphs (split on double newline)
        for para_text in text.split("\n\n"):
            stripped = para_text.strip()
            if stripped:
                doc.add_paragraph(stripped)

        # Render blocks (tables, bullet lists, etc.)
        for block in article.get("blocks", []):
            render_block(doc, block)

    # Exhibits
    exhibits = template.get("exhibits", [])
    if exhibits:
        doc.add_page_break()
        heading = doc.add_heading("Exhibits", level=1)
        heading.runs[0].font.size = Pt(12)

        doc.add_paragraph("The following Exhibits are incorporated into this Agreement:")
        doc.add_paragraph()
        for ex in exhibits:
            letter = ex.get("letter", "")
            title = ex.get("title", "")
            doc.add_paragraph(f"Exhibit {letter}: {title}", style="List Bullet")

    # Signature block
    write_signature_block(doc, scope)


def _render_section(doc, section, is_cover=False):
    """Render a section's text and blocks."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    text = section.get("text", "")
    if text:
        for para_text in text.split("\n\n"):
            stripped = para_text.strip()
            if not stripped:
                continue
            p = doc.add_paragraph(stripped)
            # Center the title line on cover pages
            if is_cover and "SUBCONTRACT AGREEMENT" in stripped:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].bold = True
                    p.runs[0].font.size = Pt(16)

    for block in section.get("blocks", []):
        render_block(doc, block)

    if text or section.get("blocks"):
        doc.add_paragraph()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Format subcontract Word document")
    parser.add_argument("--template", required=True,
                        help="JSON file with fully-populated template structure")
    parser.add_argument("--scope", required=True,
                        help="JSON file with scope/project data")
    parser.add_argument("--output", "-o", default="Subcontract.docx")
    args = parser.parse_args()
    generate_subcontract(args.template, args.scope, args.output)
